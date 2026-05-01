#!/usr/bin/env python3
"""
财报报告导出脚本 — 纯 Python 实现
Markdown → PDF / DOCX
用法：python export.py report.md --format pdf
"""

import argparse
import os
import platform
import sys
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    import markdown
except ImportError:
    sys.exit("请先安装依赖: pip install -r requirements.txt")


# ───────────────────────────── 工具函数


def find_font():
    """自动检测系统中文字体"""
    system = platform.system()
    candidates = {
        "Darwin": [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ],
        "Linux": [
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ],
        "Windows": [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        ],
    }.get(system, [])
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


# ───────────────────────────── Markdown → ElementTree 解析器


class MDParser(HTMLParser):
    """把 markdown 生成的 HTML 解析为 ElementTree，保留 text/tail 顺序"""

    def __init__(self):
        super().__init__()
        self.root = None
        self.stack = []
        self.buf = ""

    def handle_starttag(self, tag, attrs):
        self._flush()
        elem = ET.Element(tag, dict(attrs))
        if self.stack:
            self.stack[-1].append(elem)
        else:
            self.root = elem
        self.stack.append(elem)

    def handle_endtag(self, tag):
        self._flush()
        if self.stack:
            self.stack.pop()

    def handle_data(self, data):
        self.buf += data

    def _flush(self):
        if not self.buf:
            return
        if self.stack:
            elem = self.stack[-1]
            elem.text = (elem.text or "") + self.buf
        self.buf = ""

    def parse(self, md_text):
        md = markdown.Markdown(extensions=["tables", "fenced_code"])
        self.feed(md.convert(md_text))
        self._flush()
        return self.root


def iter_elem(elem):
    """按文档顺序产出事件：(text, str) | (enter, tag, attrs) | (exit, tag)"""
    if elem.text:
        yield ("text", elem.text)
    for child in elem:
        yield ("enter", child.tag, child.attrib)
        yield from iter_elem(child)
        yield ("exit", child.tag)
        if child.tail:
            yield ("text", child.tail)


# ───────────────────────────── PDF 导出器


class PDFExporter:
    def __init__(self, font_path=None):
        try:
            from fpdf import FPDF
        except ImportError:
            sys.exit("缺少 fpdf2。请运行: pip install -r requirements.txt")

        self.font_path = font_path or find_font()
        if not self.font_path:
            print("警告: 未找到中文字体，PDF 中文可能显示异常")
            print("可用 --font 手动指定，如 --font /System/Library/Fonts/PingFang.ttc")

        self.pdf = FPDF()
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.pdf.add_page()

        if self.font_path:
            try:
                self.pdf.add_font("cn", "", self.font_path)
                self.pdf.add_font("cn", "B", self.font_path)
            except Exception as e:
                print(f"字体加载失败: {e}")
                self.font_path = None

        self._f = "cn" if self.font_path else "Helvetica"

        # 状态
        self._bold = False
        self._in_table = False
        self._in_cell = False
        self._rows = []
        self._row = []
        self._cell = ""
        self._in_code = False
        self._code = []

    def _set(self, size=10, bold=False):
        style = "B" if bold else ""
        self.pdf.set_font(self._f, style, size)

    def export(self, md_text, out_path):
        self._set(10)
        for ev, *args in iter_elem(MDParser().parse(md_text)):
            if ev == "text":
                self._txt(args[0])
            elif ev == "enter":
                self._enter(args[0])
            elif ev == "exit":
                self._exit(args[0])
        self.pdf.output(str(out_path))
        print(f"PDF 已生成: {out_path}")

    def _txt(self, text):
        if self._in_code:
            self._code.append(text)
        elif self._in_cell:
            self._cell += text
        else:
            self.pdf.write(5, text)

    def _enter(self, tag):
        if tag == "h1":
            self._set(16, True)
            self.pdf.ln(8)
        elif tag == "h2":
            self._set(13, True)
            self.pdf.ln(5)
        elif tag == "h3":
            self._set(11, True)
            self.pdf.ln(4)
        elif tag == "p":
            self._set(10)
        elif tag == "blockquote":
            self._set(10)
            self.pdf.set_left_margin(self.pdf.l_margin + 8)
        elif tag == "pre":
            self._in_code = True
            self._code = []
        elif tag == "li":
            self._set(10)
            self.pdf.cell(5, 5, "·", ln=0)
        elif tag == "table":
            self._in_table = True
            self._rows = []
        elif tag == "tr":
            self._row = []
        elif tag in ("th", "td"):
            self._in_cell = True
            self._cell = ""
        elif tag == "strong":
            self._bold = True
            self._set(10, True)
        elif tag == "hr":
            y = self.pdf.get_y()
            self.pdf.line(10, y, 200, y)
            self.pdf.ln(4)

    def _exit(self, tag):
        if tag == "h1":
            self.pdf.ln(6)
        elif tag == "h2":
            self.pdf.ln(4)
        elif tag == "h3":
            self.pdf.ln(3)
        elif tag == "p":
            self.pdf.ln(4)
        elif tag == "blockquote":
            self.pdf.set_left_margin(self.pdf.l_margin - 8)
            self.pdf.ln(4)
        elif tag == "pre":
            self._in_code = False
            self._draw_code()
        elif tag == "li":
            self.pdf.ln(4)
        elif tag == "table":
            self._in_table = False
            self._draw_table()
        elif tag in ("th", "td"):
            self._in_cell = False
            self._row.append(self._cell)
        elif tag == "tr":
            self._rows.append(self._row)
        elif tag == "strong":
            self._bold = False
            self._set(10)

    def _draw_table(self):
        if not self._rows:
            return
        cols = max(len(r) for r in self._rows)
        w = (self.pdf.w - 20) / max(cols, 1)
        self._set(9)
        for row in self._rows:
            h = 5 * max(1, max(len(c) // 25 + 1 for c in row) if row else 1)
            for cell in row:
                self.pdf.cell(w, h, cell[:60], border=1, align="L")
            self.pdf.ln(h)
        self.pdf.ln(4)

    def _draw_code(self):
        text = "".join(self._code)
        if not text.strip():
            return
        self._set(8)
        self.pdf.set_fill_color(245, 245, 245)
        x = self.pdf.l_margin
        self.pdf.set_left_margin(x + 5)
        for line in text.split("\n"):
            self.pdf.cell(0, 4, line, ln=True, fill=True)
        self.pdf.set_left_margin(x)
        self.pdf.ln(2)


# ───────────────────────────── DOCX 导出器


class DOCXExporter:
    def __init__(self):
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            sys.exit("缺少 python-docx。请运行: pip install -r requirements.txt")

        self.doc = Document()
        self._set_default_font()

        self._para = None   # 当前段落
        self._run = None    # 当前 run
        self._in_table = False
        self._rows = []
        self._row = []
        self._cell = ""
        self._in_code = False

    def _set_default_font(self):
        from docx.shared import Pt
        from docx.oxml.ns import qn
        s = self.doc.styles["Normal"]
        s.font.name = "宋体"
        s.font.size = Pt(10.5)
        s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def export(self, md_text, out_path):
        for ev, *args in iter_elem(MDParser().parse(md_text)):
            if ev == "text":
                self._txt(args[0])
            elif ev == "enter":
                self._enter(args[0])
            elif ev == "exit":
                self._exit(args[0])
        self.doc.save(str(out_path))
        print(f"DOCX 已生成: {out_path}")

    def _add_run(self, text, bold=False, italic=False, size=None, font=None):
        from docx.shared import Pt
        from docx.oxml.ns import qn
        if self._para is None:
            self._para = self.doc.add_paragraph()
        r = self._para.add_run(text)
        r.font.name = font or "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font or "宋体")
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if size:
            r.font.size = Pt(size)
        self._run = r
        return r

    def _txt(self, text):
        if self._in_code:
            self._cell += text
        elif self._in_table and self._in_cell:
            self._cell += text
        else:
            self._add_run(text)

    def _enter(self, tag):
        from docx.shared import Inches, Pt

        if tag == "h1":
            self._para = self.doc.add_heading(level=1)
        elif tag == "h2":
            self._para = self.doc.add_heading(level=2)
        elif tag == "h3":
            self._para = self.doc.add_heading(level=3)
        elif tag == "p":
            self._para = self.doc.add_paragraph()
        elif tag == "blockquote":
            self._para = self.doc.add_paragraph()
            self._para.paragraph_format.left_indent = Inches(0.3)
        elif tag == "pre":
            self._in_code = True
            self._cell = ""
            self._para = self.doc.add_paragraph()
            self._para.paragraph_format.left_indent = Inches(0.3)
        elif tag == "ul" or tag == "ol":
            pass
        elif tag == "li":
            self._para = self.doc.add_paragraph(style="List Bullet")
        elif tag == "table":
            self._in_table = True
            self._rows = []
        elif tag == "tr":
            self._row = []
        elif tag in ("th", "td"):
            self._in_cell = True
            self._cell = ""
        elif tag == "strong":
            self._add_run("", bold=True)
        elif tag == "em":
            self._add_run("", italic=True)
        elif tag == "hr":
            self.doc.add_paragraph("─" * 40)

    def _exit(self, tag):
        if tag == "pre":
            self._in_code = False
            text = self._cell
            if text:
                r = self._para.add_run(text)
                r.font.name = "Courier New"
                r.font.size = Pt(9)
        elif tag == "table":
            self._in_table = False
            self._draw_table()
        elif tag in ("th", "td"):
            self._in_cell = False
            self._row.append(self._cell)
        elif tag == "tr":
            self._rows.append(self._row)

    def _draw_table(self):
        if not self._rows:
            return
        cols = max(len(r) for r in self._rows)
        table = self.doc.add_table(rows=len(self._rows), cols=cols)
        table.style = "Table Grid"
        for i, row in enumerate(self._rows):
            for j, text in enumerate(row):
                if j < cols:
                    table.rows[i].cells[j].text = text or ""


# ───────────────────────────── CLI


def main():
    parser = argparse.ArgumentParser(description="财报报告导出工具")
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("--format", choices=["pdf", "docx", "all"], default="all")
    parser.add_argument("-o", "--output", help="输出目录，默认与输入文件同级")
    parser.add_argument("--font", help="PDF 中文字体路径（默认自动检测）")
    args = parser.parse_args()

    inp = Path(args.input)
    if not inp.exists():
        sys.exit(f"错误: 文件不存在 {inp}")

    md = inp.read_text(encoding="utf-8")
    out_dir = Path(args.output) if args.output else inp.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    base = inp.stem

    if args.format in ("pdf", "all"):
        PDFExporter(font_path=args.font).export(md, out_dir / f"{base}.pdf")
    if args.format in ("docx", "all"):
        DOCXExporter().export(md, out_dir / f"{base}.docx")


if __name__ == "__main__":
    main()
